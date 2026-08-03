from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from app.docx_text import add_runs_with_markdown, set_bottom_border
from app.schemas import ResumeData

# Matches pdf_render_elegant.py's colors, which in turn match the shared LaTeX template's
# \definecolor lines (accentTitle/accentText = green, accentLine = gold).
ACCENT_GREEN = RGBColor(0x0E, 0x6E, 0x55)
ACCENT_GREEN_HEX = "0E6E55"
ACCENT_GOLD_HEX = "A16F0B"
USABLE_WIDTH = Inches(7.5)  # Letter width (8.5in) minus 0.5in left/right margins

# Times New Roman is the closest zero-install match to the LaTeX source's Cormorant
# Garamond/Charter fonts - see pdf_render_elegant.py for why those aren't embedded directly.
SERIF_FONT = "Times New Roman"


def _set_font(run, size: Pt | None = None) -> None:
    run.font.name = SERIF_FONT
    if size is not None:
        run.font.size = size


def _add_runs_with_markdown(paragraph, text: str, size: "Pt | None" = None) -> None:
    add_runs_with_markdown(paragraph, text, link_color=ACCENT_GREEN_HEX, size=size)
    for run in paragraph.runs:
        run.font.name = SERIF_FONT


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    _set_font(run, Pt(12.5))
    run.font.color.rgb = ACCENT_GREEN
    set_bottom_border(p, color=ACCENT_GOLD_HEX, size=5)


def _add_heading_with_dates(doc: Document, left_text: str, right_text: str) -> None:
    """Bold "title/school ... dates" row - the first of the two heading lines this
    template uses per job/education entry (the second, italic line carries
    company/location instead, with no date on it)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(USABLE_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    left_run = p.add_run(left_text)
    left_run.bold = True
    _set_font(left_run, Pt(10.5))
    if right_text:
        right_run = p.add_run(f"\t{right_text}")
        right_run.bold = True
        _set_font(right_run, Pt(10.5))


def _add_italic_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    _set_font(run, Pt(10.5))


def build_docx(resume: ResumeData, out_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    style = doc.styles["Normal"]
    style.font.name = SERIF_FONT
    style.font.size = Pt(10.5)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(4)
    name_run = name_p.add_run(resume.name)
    _set_font(name_run, Pt(26))
    name_run.font.color.rgb = ACCENT_GREEN
    set_bottom_border(name_p, color=ACCENT_GOLD_HEX, size=5)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(4)
    contact_p.paragraph_format.space_after = Pt(4)
    _add_runs_with_markdown(contact_p, resume.contact_line, size=Pt(9))
    set_bottom_border(contact_p, color=ACCENT_GOLD_HEX, size=5)

    _add_section_heading(doc, "Summary")
    _add_runs_with_markdown(doc.add_paragraph(), resume.summary)

    _add_section_heading(doc, "Technical Skills")
    for category, items in resume.skills.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{category}: ")
        label_run.bold = True
        _set_font(label_run, Pt(10.5))
        items_run = p.add_run(", ".join(items))
        _set_font(items_run, Pt(10.5))

    _add_section_heading(doc, "Experience")
    for job in resume.experience:
        _add_heading_with_dates(doc, job.title, job.dates)
        subheading = f"{job.company}, {job.location}" if job.location else job.company
        _add_italic_subheading(doc, subheading)
        for bullet in job.bullets:
            # En dash bullets, matching \renewcommand\labelitemi{--} in the shared
            # template - built manually (char + tab + hanging indent) rather than Word's
            # built-in "List Bullet" style, which carries hidden spacing/indent defaults
            # that survive direct paragraph_format overrides (see build_docx's own note).
            b = doc.add_paragraph()
            pf = b.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(2)
            pf.left_indent = Pt(16)
            pf.first_line_indent = Pt(-16)
            pf.tab_stops.add_tab_stop(Pt(16))
            dash_run = b.add_run("–\t")
            _set_font(dash_run, Pt(10.5))
            _add_runs_with_markdown(b, bullet)

    _add_section_heading(doc, "Education")
    for edu in resume.education:
        _add_heading_with_dates(doc, edu.school, edu.dates or "")
        subheading = f"{edu.degree}, {edu.location}" if edu.location else edu.degree
        if subheading:
            _add_italic_subheading(doc, subheading)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
