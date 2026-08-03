from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from app.docx_text import add_runs_with_markdown, set_bottom_border
from app.schemas import ResumeData

ACCENT_COLOR = RGBColor(0x0E, 0x6E, 0x55)
USABLE_WIDTH = Inches(7.5)  # Letter width (8.5in) minus 0.5in left/right margins


def _add_runs_with_markdown(paragraph, text: str, size: "Pt | None" = None) -> None:
    add_runs_with_markdown(paragraph, text, link_color="0E6E55", size=size)


def _set_bottom_border(paragraph, color: str = "0E6E55", size: int = 6) -> None:
    set_bottom_border(paragraph, color=color, size=size)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR
    _set_bottom_border(p)


def _add_heading_with_dates(doc: Document, left_text: str, right_text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(USABLE_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(left_text).bold = True
    if right_text:
        p.add_run(f"\t{right_text}")


def build_docx(resume: ResumeData, out_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(resume.name)
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ACCENT_COLOR

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2)
    _add_runs_with_markdown(contact_p, resume.contact_line, size=Pt(9.5))
    _set_bottom_border(contact_p)

    _add_section_heading(doc, "Summary")
    _add_runs_with_markdown(doc.add_paragraph(), resume.summary)

    _add_section_heading(doc, "Skills")
    for category, items in resume.skills.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{category}: ").bold = True
        p.add_run(", ".join(items))

    _add_section_heading(doc, "Experience")
    for job in resume.experience:
        _add_heading_with_dates(doc, f"{job.title} — {job.company}", job.dates)
        if job.location:
            loc_p = doc.add_paragraph()
            loc_p.paragraph_format.space_after = Pt(2)
            loc_p.paragraph_format.space_before = Pt(0)
            loc_p.add_run(job.location).italic = True
        for bullet in job.bullets:
            # Built manually (bullet char + tab + hanging indent) instead of Word's
            # built-in "List Bullet" style — that style carries hidden spacing/indent
            # defaults that survive direct paragraph_format overrides, causing a gap
            # before the first bullet that the PDF version (reportlab) doesn't have.
            b = doc.add_paragraph()
            pf = b.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(2)
            pf.left_indent = Pt(18)
            pf.first_line_indent = Pt(-18)
            pf.tab_stops.add_tab_stop(Pt(18))
            b.add_run("•\t")
            _add_runs_with_markdown(b, bullet)

    _add_section_heading(doc, "Education")
    for edu in resume.education:
        edu_heading = f"{edu.school} — {edu.degree}" if edu.degree else edu.school
        _add_heading_with_dates(doc, edu_heading, edu.dates or "")
        if edu.location:
            loc_p = doc.add_paragraph()
            loc_p.paragraph_format.space_after = Pt(2)
            loc_p.add_run(edu.location).italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
