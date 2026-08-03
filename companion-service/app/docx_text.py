from __future__ import annotations

import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

_TOKEN_RE = re.compile(r"\*\*(.+?)\*\*|\[([^\]]+)\]\(([^)]+)\)")


def add_hyperlink(paragraph, url: str, text: str, color: str = "0E6E55", size: "Pt | None" = None) -> None:
    """python-docx has no built-in hyperlink support - a real docx hyperlink is a
    <w:hyperlink> element wrapping a run, pointing at an external relationship, which
    has to be built by hand via the OOXML APIs. Shared by every DOCX template."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rPr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))  # w:sz is in half-points
        rPr.append(sz)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_runs_with_markdown(paragraph, text: str, link_color: str = "0E6E55", size: "Pt | None" = None) -> None:
    """Split the LLM's markdown-style **bold** spans and [text](url) links into
    alternating bold/link/plain runs. (No XML-escaping needed for plain/bold runs -
    unlike reportlab, python-docx runs store plain text and escape internally when
    serializing, so a literal "&" or "<" is never an issue.)"""
    pos = 0
    for match in _TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            if size is not None:
                run.font.size = size
        if match.group(1) is not None:
            run = paragraph.add_run(match.group(1))
            run.bold = True
            if size is not None:
                run.font.size = size
        else:
            add_hyperlink(paragraph, match.group(3), match.group(2), color=link_color, size=size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if size is not None:
            run.font.size = size


def set_bottom_border(paragraph, color: str = "0E6E55", size: int = 6) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
